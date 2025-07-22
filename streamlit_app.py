import streamlit as st
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import datetime
from docx.text.run import Run
from docx.shared import Inches, RGBColor
import os
from dotenv import load_dotenv

# ──────────────────────────────────────────────
# Table handling
# ──────────────────────────────────────────────
def extract_tables(wb):
# Table 1
    ws1 = wb["Table 1"]
    rows1 = list(ws1.iter_rows(min_row=2, values_only=True))
    rows1 = [r for r in rows1 if r and r[0] and r[1]]
    header1 = list(rows1[0][:-1])
    body1 = []
    for r in rows1[1:]:
        row_cleaned = []
        for i, val in enumerate(r[:-1]):
            if isinstance(val, (datetime.datetime, datetime.date)):
                val = val.strftime("%B %d, %Y")
            if i == 1 and isinstance(val, str):  # Proper-case Name of Shareholder
                val = val.title()
            row_cleaned.append(str(val) if val is not None else "")
        body1.append(row_cleaned)
    table1_data = [header1] + body1

    # Table 2
    ws2 = wb["Table 2"]
    rows2 = list(ws2.iter_rows(min_row=2, values_only=True))
    rows2 = [r for r in rows2 if r and r[0] and r[1]]
    header2 = list(rows2[0][:-1])
    body2 = []
    for r in rows2[1:]:
        row_cleaned = []
        for i, val in enumerate(r[:-1]):
            if isinstance(val, (datetime.datetime, datetime.date)):
                val = val.strftime("%B %d, %Y")
            if i == 1 and isinstance(val, str):
                val = val.title()
            row_cleaned.append(str(val) if val is not None else "")
        body2.append(row_cleaned)
    table2_data = [header2] + body2

    # Table 5 - V1
    # Table 5
    ws5 = wb["Table 5"]
    header5 = [c.value for c in ws5[2]]

    # Columns we want
    wanted_cols = ["Sr. No.", "Name of Employee", "Designation", "Area of Expertise"]
    col_indices = [i for i, h in enumerate(header5) if h in wanted_cols]
    rows5 = list(ws5.iter_rows(min_row=3, values_only=True))
    rows5 = [r for r in rows5 if r and r[0] and r[1]]

    # Table 5 V1 – cleaned, proper-case name
    table5_v1_data = [[header5[i] for i in col_indices]]
    for r in rows5:
        row_cleaned = []
        for j, i in enumerate(col_indices):
            val = r[i]
            if isinstance(val, (datetime.datetime, datetime.date)):
                val = val.strftime("%B %d, %Y")
            elif header5[i] == "Name of Employee" and isinstance(val, str):
                val = val.title()
            row_cleaned.append(str(val).strip() if val else "")
        table5_v1_data.append(row_cleaned)

    # Table 5 V2 – specialization table
    col_area = header5.index("Area of Expertise")
    col_name = header5.index("Name of Employee")
    table5_v2_data = [["Specialization", "Name of employee"]]
    for r in rows5:
        area = r[col_area]
        name = r[col_name]
        if name and isinstance(name, str):
            name = name.title()
        table5_v2_data.append([
            str(area).strip() if area else "",
            str(name).strip() if name else ""
        ])

    return {
        "table1_data": table1_data,
        "table2_data": table2_data,
        "table5_v1_data": table5_v1_data,
        "table5_v2_data": table5_v2_data
    }

# ──────────────────────────────────────────────
# Word table insertion utils
# ──────────────────────────────────────────────
def add_grid(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)

def insert_table_as_paragraph(paragraph, data):
    doc = paragraph._parent.part.document
    parent = paragraph._p.getparent()
    idx = parent.index(paragraph._p)
    parent.remove(paragraph._p)

    tbl = doc.add_table(rows=0, cols=len(data[0]))
    tbl.style = 'Table Grid'
    add_grid(tbl)

    # Apply fixed column widths for Table 1 (only if 4 cols)
    if len(data[0]) == 4:
        widths = [0.8, 2.4, 2.4, 2.4]  # ~10%, 30%, 30%, 30%
        for i, width in enumerate(widths):
            tbl.columns[i].width = Inches(width)

    hdr = tbl.add_row().cells
    for i, val in enumerate(data[0]):
        run = hdr[i].paragraphs[0].add_run(val)
        run.bold = True
        run.font.size = Pt(9)

    for row in data[1:]:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    parent.insert(idx, tbl._tbl)

def insert_table_in_cell(cell, data):
    cell.text = ""
    tbl = cell.add_table(rows=0, cols=len(data[0]))
    add_grid(tbl)

    hdr = tbl.add_row().cells
    for i, val in enumerate(data[0]):
        run = hdr[i].paragraphs[0].add_run(val)
        run.bold = True
        run.font.size = Pt(9)

    for row in data[1:]:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

def apply_final_styling(doc):
    # Apply Arial 12 pt to all paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Headers
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Tables – Arial 11 pt
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)

def replace_placeholder_in_paragraph(paragraph, mapping):
    full_text = paragraph.text
    for k, v in mapping.items():
        full_text = full_text.replace(f"${{{k}}}", v)
    # Clear all existing runs
    for run in paragraph.runs:
        run.text = ""
    # Add the replaced full text as one run
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)




# ──────────────────────────────────────────────
# Password Popup Logic
# ──────────────────────────────────────────────
load_dotenv()
APP_PASSWORD = os.getenv("APP_PASSWORD", "")

if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

def password_popup():
    # Render overlay background only
    st.markdown("""
    <style>
    .password-popup-bg {
        position: absolute;
        top: 0; left: 0; width: 100vw; height: 100vh;
        background: rgba(0,0,0,0.5);
        z-index: 1;
        pointer-events: none;
    }
    .password-box {
        position: relative;
        z-index: 2;
    }
    </style>
    <div class="password-popup-bg"></div>
    """, unsafe_allow_html=True)
    # Render password box in a Streamlit block (always interactive)
    box = st.empty()
    with box.container():
        st.markdown("<div class='password-box' style='margin:auto; max-width:400px;'>", unsafe_allow_html=True)
        st.markdown("#### 🔒 Enter Password to Access the App")
        pw1, pw2, pw3, pw4 = st.columns(4)
        pw_vals = []
        for i, col in enumerate([pw1, pw2, pw3, pw4]):
            pw_vals.append(col.text_input(f"Password digit {i+1}", type="password", key=f"pw_{i}", max_chars=1, label_visibility="collapsed"))
        pw_input = "".join(pw_vals)
        submit = st.button("Submit", key="pw_submit")
        if submit:
            if pw_input == APP_PASSWORD:
                st.session_state["authenticated"] = True
            else:
                st.error("Incorrect password. Please try again.")
        st.markdown("</div>", unsafe_allow_html=True)

if not st.session_state["authenticated"]:
    password_popup()
    st.stop()

# ──────────────────────────────────────────────
# Main Streamlit UI
# ──────────────────────────────────────────────
st.set_page_config("ERP Application Checker", layout="wide")
st.markdown("""
<style>
.scroll-box {
    max-height: 400px;
    overflow-y: auto;
    padding: 0.5rem;
    border: 1px solid #ccc;
    background-color: #f9f9f9; /* Slightly different background for visibility */
    margin-bottom: 1rem; /* Add some space below the box */
}
</style>
""", unsafe_allow_html=True)

st.title("📄 ERP Category II – Compliance Check + Office Note Generator")

uploaded_file = st.file_uploader("📑 Upload filled *Excel Format.xlsx*", type=["xlsx"], label_visibility="hidden")

if uploaded_file:
    try:
        wb = openpyxl.load_workbook(uploaded_file, data_only=True)
        analysis_sheet = wb["Analysis"]

        # Compliance display
        all_errors = []
        for row in analysis_sheet.iter_rows(min_row=2, min_col=1, max_col=1):
            val = row[0].value
            if val and isinstance(val, str) and val.strip():
                all_errors.append(val.strip())

        st.subheader(f"🔍 Deficiencies and Non-compliances found ({len(all_errors)})")
        if all_errors:
            # Build HTML for each error, mimicking st.error styling
            error_items_html = ""
            for err in all_errors:
                 # Basic mimic of st.error styling
                 # Ref: Based on inspecting typical st.error output
                 error_html = f'<div style="background-color: #fff0f0; border: 1px solid #ff4b4b; border-left: 5px solid #ff4b4b; border-radius: 0.25rem; padding: 1rem; margin-bottom: 0.5rem; color: #31333F; word-wrap: break-word;">{err}</div>'
                 error_items_html += error_html

            # Wrap errors in the scroll-box div (using the class defined in the <style> tag)
            scroll_box_html = f'<div class="scroll-box">{error_items_html}</div>'
            # Render the whole block as raw HTML
            st.markdown(scroll_box_html, unsafe_allow_html=True)
        else:
            st.success("✅ No compliance issues detected!")

        # Office Note generation
        buffer = None

        
        try:
            doc = Document("SampleOfficeNote.docx")

            def extract_named(sheet, col_var=3, col_val=4):
                values = {}
                for row in sheet.iter_rows(min_row=2):
                    var = row[col_var].value
                    val = row[col_val].value
                    if not var:
                        continue
                    
                    var_name = str(var).strip()
                    if isinstance(val, datetime.date):
                        val = val.strftime("%B %d, %Y")
                    elif var_name in {
                        "applicant_name", "regd_address", "corr_address",
                        "comp_officer_name", "cont_person_name"
                    } and isinstance(val, str):
                        val = val.title()
                    values[var_name] = str(val).strip() if val else ""

                return values

            vars_basic = extract_named(wb["Basic Details"], 2, 3)
            vars_elig = extract_named(wb["Eligibility Criteria"], 3, 4)
            vars_all = {**vars_basic, **vars_elig}

            if "applicant_name" in vars_all:
                abbr = ''.join(w[0] for w in vars_all["applicant_name"].split()).upper()
                vars_all["applicant_name_abb"] = abbr

            # MD / CEO processing
            md = vars_all.get("md_name", "").strip().lower()
            ceo = vars_all.get("ceo_name", "").strip().lower()
            md_val = "" if md in ["", "na"] else md
            ceo_val = "" if ceo in ["", "na"] else ceo

            if not md_val and not ceo_val:
                vars_all["md_ceo_name"] = ""
                vars_all["md_ceo_rating"] = ""
            else:
                vars_all["md_ceo_name"] = f"Name of MD: {md_val or 'NA'}\nName of CEO: {ceo_val or 'NA'}"
                entity = ("the MD" if md_val and not ceo_val else
                        "the CEO" if ceo_val and not md_val else
                        "MD & CEO" if md_val == ceo_val else
                        "both the MD and CEO")
                name_str = md_val if md_val == ceo_val else f"{md_val} and {ceo_val}"
                vars_all["md_ceo_rating"] = f"The Applicant has submitted that {entity} {name_str} is/are not part of rating decisions by the ERP."

            # Phrase replacements for sentence-friendly language
            # ── Preprocess logic-based variable values ──
            phrase_map = {}

            # whether_declaration
            val = vars_all.get("whether_declaration", "").strip().lower()
            phrase_map["whether_declaration"] = "submitted" if val == "yes" else "NOT submitted"

            # whether_revenue_clients
            val = vars_all.get("whether_revenue_clients", "").strip().lower()
            if val == "yes":
                phrase_map["whether_revenue_clients"] = (
                    "Yes. The Business Plan contains information about target revenue and the targeted number of clients it plans to service, within 2 years of obtaining a certificate."
                )
            elif val == "no":
                phrase_map["whether_revenue_clients"] = (
                    "No. The Business Plan DOES NOT contain information about target revenue and the targeted number of clients."
                )

            # whether_breakeven_date
            val = vars_all.get("whether_breakeven_date", "").strip().lower()
            if val == "yes":
                phrase_map["whether_breakeven_date"] = (
                    "Yes. The Business Plan contains information about Target Breakeven Date."
                )
            elif val == "no":
                phrase_map["whether_breakeven_date"] = (
                    "No. The Business Plan DOES NOT contain information about Target Breakeven Date."
                )

            # whether_cash_losses
            val = vars_all.get("whether_cash_losses", "").strip().lower()
            if val == "yes":
                phrase_map["whether_cash_losses"] = (
                    "Yes. The Business Plan contains information about the cumulative cash losses that the applicant projects to incur until the targeted breakeven date, along with the activities or areas in which such losses shall be incurred."
                )
            elif val == "no":
                phrase_map["whether_cash_losses"] = (
                    "No. The Business Plan DOES NOT contain the required information on projected cumulative losses."
                )

            # Fallback for business_plan_summary
            val = vars_all.get("business_plan_summary", "").strip()
            if not val:
                phrase_map["business_plan_summary"] = "The Applicant has not submitted a summary of its Business Plan."
            else:
                phrase_map["business_plan_summary"] = val

            # Infrastructure + Operations Handling
            ops = vars_all.get("operations", "").strip().lower()
            undertaking = vars_all.get("operations_undertaking", "").strip().lower()

            infra_yes = "has submitted an undertaking that it has "
            infra_no  = "has NOT SUBMITTED an undertaking that it has "

            if "Only remote" in ops:
                infra_phrase = "necessary infrastructure including technology, equipment and manpower, to enable it to provide ESG rating services."
            else:
                infra_phrase = "necessary infrastructure including adequate office space, technology, equipment and manpower, to enable it to provide ESG rating services."

            if undertaking == "undertaking provided":
                phrase_map["operations_undertaking"] = infra_yes + infra_phrase
            else:
                phrase_map["operations_undertaking"] = infra_no + infra_phrase


            # Combine everything
            vars_all.update(phrase_map)

            # Extract tables
            tables = extract_tables(wb)

            # Replace in paragraphs
            for p in doc.paragraphs:
                replace_placeholder_in_paragraph(p, vars_all)

            for section in doc.sections:
                for p in section.header.paragraphs:
                    replace_placeholder_in_paragraph(p, vars_all)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if any(token in p.text for token in ["${table5_v1_data}", "${table5_v2_data}"]):
                                continue  # skip for now – will be replaced by insert_table_in_cell()
                            replace_placeholder_in_paragraph(p, vars_all)

                        # Now handle the nested tables
                        for key, data in {
                            "table5_v1_data": tables["table5_v1_data"],
                            "table5_v2_data": tables["table5_v2_data"]
                        }.items():
                            if f"${{{key}}}" in cell.text:
                                insert_table_in_cell(cell, data)


            # Insert paragraph-based tables
            for para in doc.paragraphs:
                for key, data in {"table1_data": tables["table1_data"], "table2_data": tables["table2_data"]}.items():
                    if f"${{{key}}}" in para.text:
                        para.text = ""
                        insert_table_as_paragraph(para, data)
            
            apply_final_styling(doc)
        
            # Save and download
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.success("✅ Office Note generated successfully.")
        except Exception as e:
            st.error(f"Error generating Office Note: {e}")
        
        if buffer:
            st.download_button(
                "⬇ Download Office Note (.docx)",
                buffer,
                file_name="OfficeNote.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"Error reading Excel file: {e}")
else:
    st.info("Please upload the Excel file.")
